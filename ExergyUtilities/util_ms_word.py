# ExergyUtilities
# Copyright (c) 2010, B. Marcus Jones <>
# All rights reserved.
#
#    This program is free software: you can redistribute it and/or modify
#    it under the terms of the GNU General Public License as published by
#    the Free Software Foundation, either version 3 of the License, or
#    (at your option) any later version.
#
#    This program is distributed in the hope that it will be useful,
#    but WITHOUT ANY WARRANTY; without even the implied warranty of
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#    GNU General Public License for more details.
#
#    You should have received a copy of the GNU General Public License
#    along with this program.  If not, see <http://www.gnu.org/licenses/>.

"""The :mod:`xx` module is a utility and notebook module to interface with Excel. The two xlrd/xlwt  
"""

from win32com.client import Dispatch
import logging.config
import os
import re
import unittest
#from config import *
from docx.shared import Inches
from docx import Document
import docx as docx_import
from docx.enum.style import WD_STYLE_TYPE
import shutil
from util_pretty_print import print_table


#path_folder = r"C:\Users\Admin\Desktop\\"
path_folder = r"C:\CesCloud Senegal PV\10 Tender Dossier\03 Dossier\\"

path_folder = r"C:\Users\jon\Desktop"

file_name = r"test.docx"
file_name = r"MASTER r12.docx"
file_name = r"P02 S05 Cahier des Clauses Techniques et Services (CCTS) r15.docx"


full_path = os.path.join(path_folder,file_name)


def test1():
    print("Hello")
    
    doc = Document()
    print(doc)
    
    document.add_heading('Document Title', 0)
    
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')
    
    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )
    
    #document.add_picture('monty-truth.png', width=Inches(1.25))
    
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
#     for item in recordset:
#         row_cells = table.add_row().cells
#         row_cells[0].text = str(item.qty)
#         row_cells[1].text = str(item.id)
#         row_cells[2].text = item.desc
    
    document.add_page_break()
    print(full_path)
    #raise
    document.save(full_path)

def replace_links():
    doc = Document(full_path)
    print("***PARAGRAPHS")
    for para in doc.paragraphs:
        #print(para, para.style.name)
        if para.style.name == "00 LINK":
            for run in para.runs:
                pass
                #print("\t",run, run.style.name)
                #print(run.style.name,repr(run.text))
                
            print("".join([run.text for run in para.runs]))

def list_para_styles(doc):

    builtin_paragraph_styles = [s for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH and s.builtin == True]
    not_builtin_paragraph_styles = [s for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH and s.builtin == False]

    print("***BUILTIN PARA STYLES***")
    header= ["i", "Name", "Object", "Base", "Hidden"]
    separate = ["-" * len(head) for head in header]
    rows = list()
    rows.append(header)
    rows.append(separate)
    for i,style in enumerate(builtin_paragraph_styles):
        try:
            base = str(style.base_style.name)
        except: 
            base = "***None***"
        row = [i, style.name, str(style), base, style.hidden]
        rows.append(row)

    print_table(rows)

    print("***CUSTOM PARA STYLES***")
    header= ["i", "Name", "Object", "Base", "Hidden"]
    separate = ["-" * len(head) for head in header]
    rows = list()
    rows.append(header)
    rows.append(separate)
    for i,style in enumerate(not_builtin_paragraph_styles):
        try:
            base = str(style.base_style.name)
        except: 
            base = "***None***"
        row = [i, style.name, str(style), base, style.hidden]
        rows.append(row)

    print_table(rows)

           
def return_para_styles(doc):
    return [s for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
    

           
def list_styles():
    doc = Document(full_path)
    print(doc)    
    
    print("***STYLES***")
    for i,style in enumerate(doc.styles):
        try:
            base = style.base_style
        except: 
            base = "None"
        print("{:5} {:50} {:70} {:20} {}".format(i,str(type(style)),str(style),style.name,base))
        #print("{:5}|{:30}|{:30}|{:30}".format(i,type(style),style,style.name))
    
def test2():
    doc = Document(full_path)
    print(doc)
    print("***SECTIONS***")
    for sect in doc.sections:
        #print(sect.name ,)
        print(sect.start_type,sect.orientation, sect.page_width, sect.page_height)


    paragraph_styles = [s for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
    
    print("***PARA STYLES***")
    for style in paragraph_styles:
        print(style, style.name)
            
    print(style.font.name, style.font.size)
    
    print("***PARAGRAPHS")
    for para in doc.paragraphs:
        print(para, para.style.name)
        for run in para.runs:
            print("\t",run, run.style.name)
            print(run.text)

def list_paras(doc):
    print("***PARAGRAPHS")
    for i,para in enumerate(doc.paragraphs):
        print(i,para, para.style.name)
        for run in para.runs:
            print("\t {:30} {}".format(run.style.name, run.text))
            #print(run.text)


        
def main():
    doc = Document(full_path)
    print(doc)
    list_paras(doc)
    #list_para_styles(doc)
    #replace_links()
    
if __name__ == "__main__":
    main()
    